from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import logging
import re
import asyncio
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Dict, Optional, Union
from enum import Enum
import uuid
from datetime import datetime
import io
from docx import Document
import tempfile
import json

# Try to load spaCy
nlp = None
try:
    import spacy
    nlp = spacy.load("fr_core_news_lg")
    print("✅ spaCy French model loaded successfully")
except Exception as e:
    print(f"⚠️ spaCy French model not found: {e}")
    nlp = None

# Try to load httpx for Ollama
httpx = None
try:
    import httpx
    print("✅ httpx loaded for Ollama integration")
except ImportError:
    print("⚠️ httpx not found - Ollama integration disabled")

# Create the main app
app = FastAPI(title="Anonymiseur Juridique RGPD v3.0")
api_router = APIRouter(prefix="/api")

# Models
class ProcessingMode(str, Enum):
    STANDARD = "standard"
    ADVANCED = "advanced" 
    OLLAMA = "ollama"

class EntityType(str, Enum):
    PHONE = "phone"
    EMAIL = "email"
    SIRET = "siret"
    SSN = "ssn"
    ADDRESS = "address"
    LEGAL = "legal"
    PERSON = "person"
    ORGANIZATION = "organization"

class EntitySource(str, Enum):
    REGEX = "REGEX"
    NER = "NER"
    OLLAMA = "OLLAMA"
    MANUAL = "MANUAL"

class Position(BaseModel):
    start: int
    end: int

class Entity(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    text: str
    type: EntityType
    source: EntitySource
    confidence: float
    replacement: str
    positions: List[Position]
    selected: bool = True

class OllamaConfig(BaseModel):
    url: str = "http://localhost:11434"
    model: str = "llama3.2:3b"
    custom_prompt: Optional[str] = None
    timeout: int = 30

class DocumentRequest(BaseModel):
    content: str
    filename: str
    mode: ProcessingMode
    ollama_config: Optional[OllamaConfig] = None

class ProcessingResponse(BaseModel):
    entities: List[Entity]
    processing_time: float
    mode_used: ProcessingMode
    total_occurrences: int
    spacy_available: bool
    ollama_available: bool

class DocumentGenerationRequest(BaseModel):
    entities: List[Entity]
    original_content: str
    filename: str = "document_anonymise.docx"

# Services
class RegexService:
    def __init__(self):
        # CORRECTED: Patterns téléphones français - SIMPLIFIÉS et TESTÉS
        self.phone_patterns = [
            r'\b0[1-9](?:[\s\.\-]?\d{2}){4}\b',  # 06.12.34.56.78, 06 12 34 56 78, etc.
            r'\+33[\s\.\-]?[1-9](?:[\s\.\-]?\d{2}){4}\b', # +33 6 12 34 56 78
        ]
        
        # CORRECTED: Email pattern - Plus robuste
        self.email_pattern = r'\b[a-zA-Z0-9][a-zA-Z0-9._%+-]*@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
        
        # CORRECTED: SIRET pattern - Exactement 14 chiffres
        self.siret_pattern = r'\b\d{14}\b'
        
        # CORRECTED: French SSN pattern - Plus flexible
        self.ssn_pattern = r'\b[12][\s\-]?\d{2}[\s\-]?\d{2}[\s\-]?\d{2}[\s\-]?\d{3}[\s\-]?\d{3}[\s\-]?\d{2}\b'
        
        # CORRECTED: Adresses françaises - Patterns améliorés
        self.address_patterns = [
            r'\b\d+[\s,]*(?:bis|ter)?[\s,]+(?:rue|avenue|av|boulevard|bd|place|pl|impasse|imp|allée|chemin|route|rt)[\s,]+[a-zA-Z\s\-\']+\b',
            r'\b\d{5}[\s,]+[A-Z][a-zA-Z\s\-\']{2,}\b',  # Code postal + ville (min 3 chars)
        ]
        
        # CORRECTED: Références juridiques - Plus spécifiques
        self.legal_patterns = [
            r'\bRG[\s\-]?\d+[\s\-/]\d+\b',
            r'\bdossier[\s]+(?:n°?|numéro)[\s]*\d+[\s\-/]\d+\b',
            r'\barticle[\s]+\d+(?:[\s\-]\d+)?\b',
        ]

    def simple_luhn_check(self, number: str) -> bool:
        """Validation Luhn simplifiée pour SIRET"""
        if len(number) != 14 or not number.isdigit():
            return False
        
        # Calcul Luhn basique
        total = 0
        reverse_digits = number[::-1]
        
        for i, digit in enumerate(reverse_digits):
            n = int(digit)
            if i % 2 == 1:  # Every second digit
                n *= 2
                if n > 9:
                    n = n // 10 + n % 10
            total += n
        
        return total % 10 == 0

    def extract_entities(self, text: str) -> List[Entity]:
        entities = []
        
        print(f"🔍 REGEX - Analysing text: '{text[:100]}...'")
        
        # Téléphones
        for i, pattern in enumerate(self.phone_patterns):
            try:
                matches = list(re.finditer(pattern, text, re.IGNORECASE))
                print(f"📞 Pattern {i+1} '{pattern}' found {len(matches)} matches")
                for match in matches:
                    phone_text = match.group().strip()
                    print(f"📞 Found phone: '{phone_text}' at {match.start()}-{match.end()}")
                    entities.append(Entity(
                        text=phone_text,
                        type=EntityType.PHONE,
                        source=EntitySource.REGEX,
                        confidence=1.0,
                        replacement="06 XX XX XX XX",
                        positions=[Position(start=match.start(), end=match.end())]
                    ))
            except Exception as e:
                print(f"❌ Error in phone pattern {i}: {e}")
        
        # Emails
        try:
            matches = list(re.finditer(self.email_pattern, text, re.IGNORECASE))
            print(f"📧 Email pattern found {len(matches)} matches")
            for match in matches:
                email_text = match.group().strip()
                print(f"📧 Found email: '{email_text}' at {match.start()}-{match.end()}")
                entities.append(Entity(
                    text=email_text,
                    type=EntityType.EMAIL,
                    source=EntitySource.REGEX,
                    confidence=1.0,
                    replacement="email.anonymise@exemple.fr",
                    positions=[Position(start=match.start(), end=match.end())]
                ))
        except Exception as e:
            print(f"❌ Error in email pattern: {e}")
        
        # SIRET avec validation Luhn
        try:
            matches = list(re.finditer(self.siret_pattern, text))
            print(f"🏭 SIRET pattern found {len(matches)} potential matches")
            for match in matches:
                siret = match.group().strip()
                print(f"🏭 Checking SIRET: '{siret}' at {match.start()}-{match.end()}")
                if self.simple_luhn_check(siret):
                    print(f"✅ Valid SIRET: {siret}")
                    entities.append(Entity(
                        text=siret,
                        type=EntityType.SIRET,
                        source=EntitySource.REGEX,
                        confidence=1.0,
                        replacement="[SIRET Anonymisé]",
                        positions=[Position(start=match.start(), end=match.end())]
                    ))
                else:
                    print(f"❌ Invalid SIRET (Luhn): {siret}")
        except Exception as e:
            print(f"❌ Error in SIRET pattern: {e}")
        
        # SSN
        try:
            matches = list(re.finditer(self.ssn_pattern, text))
            print(f"🆔 SSN pattern found {len(matches)} matches")
            for match in matches:
                ssn_text = match.group().strip()
                print(f"🆔 Found SSN: '{ssn_text}' at {match.start()}-{match.end()}")
                entities.append(Entity(
                    text=ssn_text,
                    type=EntityType.SSN,
                    source=EntitySource.REGEX,
                    confidence=1.0,
                    replacement="[N° Sécurité Sociale Anonymisé]",
                    positions=[Position(start=match.start(), end=match.end())]
                ))
        except Exception as e:
            print(f"❌ Error in SSN pattern: {e}")
        
        # Adresses
        for i, pattern in enumerate(self.address_patterns):
            try:
                matches = list(re.finditer(pattern, text, re.IGNORECASE))
                print(f"🏠 Address pattern {i+1} found {len(matches)} matches")
                for match in matches:
                    address_text = match.group().strip()
                    # Filtrer les faux positifs (trop courts)
                    if len(address_text) > 8:
                        print(f"🏠 Found address: '{address_text}' at {match.start()}-{match.end()}")
                        entities.append(Entity(
                            text=address_text,
                            type=EntityType.ADDRESS,
                            source=EntitySource.REGEX,
                            confidence=1.0,
                            replacement="[Adresse Anonymisée]",
                            positions=[Position(start=match.start(), end=match.end())]
                        ))
            except Exception as e:
                print(f"❌ Error in address pattern {i}: {e}")
        
        # Références juridiques
        for i, pattern in enumerate(self.legal_patterns):
            try:
                matches = list(re.finditer(pattern, text, re.IGNORECASE))
                print(f"⚖️ Legal pattern {i+1} found {len(matches)} matches")
                for match in matches:
                    legal_text = match.group().strip()
                    print(f"⚖️ Found legal ref: '{legal_text}' at {match.start()}-{match.end()}")
                    entities.append(Entity(
                        text=legal_text,
                        type=EntityType.LEGAL,
                        source=EntitySource.REGEX,
                        confidence=1.0,
                        replacement="[Référence Anonymisée]",
                        positions=[Position(start=match.start(), end=match.end())]
                    ))
            except Exception as e:
                print(f"❌ Error in legal pattern {i}: {e}")
        
        print(f"✅ REGEX - Total entities found: {len(entities)}")
        return entities

class NERService:
    def __init__(self):
        self.nlp = nlp
        self.person_counter = 1
        self.org_counter = 1
    
    def extract_entities(self, text: str) -> List[Entity]:
        if not self.nlp:
            print("⚠️ NER - spaCy not available")
            return []
        
        print(f"🧠 NER - Processing text with spaCy...")
        entities = []
        
        try:
            doc = self.nlp(text)
            
            for ent in doc.ents:
                print(f"🔍 NER - Found: '{ent.text}' ({ent.label_}) at {ent.start_char}-{ent.end_char}")
                
                if ent.label_ in ["PER", "PERSON"]:  # Person
                    entities.append(Entity(
                        text=ent.text.strip(),
                        type=EntityType.PERSON,
                        source=EntitySource.NER,
                        confidence=0.9,
                        replacement=f"Personne {chr(64 + self.person_counter)}",
                        positions=[Position(start=ent.start_char, end=ent.end_char)]
                    ))
                    self.person_counter += 1
                
                elif ent.label_ in ["ORG", "ORGANIZATION"]:  # Organization
                    entities.append(Entity(
                        text=ent.text.strip(),
                        type=EntityType.ORGANIZATION,
                        source=EntitySource.NER,
                        confidence=0.85,
                        replacement=f"Organisation {chr(64 + self.org_counter)}",
                        positions=[Position(start=ent.start_char, end=ent.end_char)]
                    ))
                    self.org_counter += 1
        
        except Exception as e:
            print(f"❌ NER Error: {e}")
            return []
        
        print(f"✅ NER - Total entities found: {len(entities)}")
        return entities

class OllamaService:
    def __init__(self, config: OllamaConfig):
        self.config = config
    
    async def check_availability(self) -> bool:
        if not httpx:
            return False
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(f"{self.config.url}/api/tags", timeout=3.0)
                return response.status_code == 200
        except:
            return False
    
    async def get_available_models(self) -> List[str]:
        if not httpx:
            return []
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(f"{self.config.url}/api/tags", timeout=5.0)
                if response.status_code == 200:
                    data = response.json()
                    return [model["name"] for model in data.get("models", [])]
        except:
            pass
        return []
    
    async def extract_entities(self, text: str) -> List[Entity]:
        """Extract entities using Ollama (placeholder for now)"""
        return []

# Document processor
class DocumentProcessor:
    @staticmethod
    def apply_anonymization(content: str, entities: List[Entity]) -> str:
        """Apply entity replacements to text content"""
        # Sort by position (start) in descending order to avoid position shifts
        sorted_entities = sorted(
            [e for e in entities if e.selected], 
            key=lambda x: x.positions[0].start, 
            reverse=True
        )
        
        result = content
        for entity in sorted_entities:
            pos = entity.positions[0]
            result = result[:pos.start] + entity.replacement + result[pos.end:]
        
        return result

# Initialize services
regex_service = RegexService()
ner_service = NERService()

# API Endpoints
@api_router.get("/")
async def root():
    return {"message": "Anonymiseur Juridique RGPD v3.0", "status": "running"}

@api_router.get("/health")
async def health_check():
    """Check system status"""
    return {
        "status": "healthy",
        "spacy_available": nlp is not None,
        "ollama_available": False,
        "timestamp": datetime.now().isoformat()
    }

@api_router.post("/process", response_model=ProcessingResponse)
async def process_document(request: DocumentRequest):
    """Process document with selected mode"""
    print(f"🚀 [API] Processing request:")
    print(f"   - Filename: {request.filename}")
    print(f"   - Mode: {request.mode}")
    print(f"   - Content length: {len(request.content)} chars")
    print(f"   - Content preview: {request.content[:200]}...")
    
    start_time = datetime.now()
    all_entities = []
    
    try:
        # Always run REGEX
        print(f"🔧 Running REGEX processing...")
        regex_entities = regex_service.extract_entities(request.content)
        all_entities.extend(regex_entities)
        print(f"📊 REGEX found {len(regex_entities)} entities")
        
        # Run NER if Advanced mode and spaCy available
        if request.mode == ProcessingMode.ADVANCED and nlp:
            print(f"🔧 Running NER processing...")
            ner_entities = ner_service.extract_entities(request.content)
            all_entities.extend(ner_entities)
            print(f"📊 NER found {len(ner_entities)} entities")
        elif request.mode == ProcessingMode.ADVANCED and not nlp:
            print(f"⚠️ Advanced mode requested but spaCy not available")
        
        # Ollama mode (placeholder for now)
        if request.mode == ProcessingMode.OLLAMA:
            print(f"⚠️ Ollama mode not implemented yet")
        
        # Remove overlapping entities (keep highest confidence)
        unique_entities = []
        for entity in all_entities:
            # Check for overlaps
            overlaps = False
            for existing in unique_entities[:]:
                if (entity.positions[0].start < existing.positions[0].end and 
                    entity.positions[0].end > existing.positions[0].start):
                    # Overlap detected - keep the one with higher confidence
                    if entity.confidence > existing.confidence:
                        unique_entities.remove(existing)
                        print(f"🔄 Replaced lower confidence entity: '{existing.text}' with '{entity.text}'")
                    else:
                        overlaps = True
                        print(f"🔄 Skipped lower confidence entity: '{entity.text}'")
                        break
            
            if not overlaps:
                unique_entities.append(entity)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        print(f"✅ Processing completed successfully:")
        print(f"   - Total entities: {len(unique_entities)}")
        print(f"   - Processing time: {processing_time:.2f}s")
        print(f"   - Mode used: {request.mode}")
        
        response = ProcessingResponse(
            entities=unique_entities,
            processing_time=processing_time,
            mode_used=request.mode,
            total_occurrences=len(unique_entities),
            spacy_available=nlp is not None,
            ollama_available=False
        )
        
        print(f"📤 Sending response with {len(response.entities)} entities")
        return response
        
    except Exception as e:
        print(f"❌ Error during processing: {str(e)}")
        print(f"❌ Error type: {type(e).__name__}")
        import traceback
        print(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")

@api_router.post("/test-ollama")
async def test_ollama_connection(config: OllamaConfig):
    """Test Ollama connection"""
    service = OllamaService(config)
    available = await service.check_availability()
    models = await service.get_available_models() if available else []
    
    return {
        "connected": available,
        "models": models,
        "config": config.dict()
    }

@api_router.get("/ollama-models")
async def get_ollama_models(url: str = "http://localhost:11434"):
    """Get available Ollama models"""
    config = OllamaConfig(url=url)
    service = OllamaService(config)
    models = await service.get_available_models()
    return {"models": models}

@api_router.post("/generate-document")
async def generate_anonymized_document(request: DocumentGenerationRequest):
    """Generate anonymized DOCX document"""
    try:
        print(f"📄 [API] Generating document:")
        print(f"   - Filename: {request.filename}")
        print(f"   - Entities count: {len(request.entities)}")
        print(f"   - Selected entities: {len([e for e in request.entities if e.selected])}")
        print(f"   - Content length: {len(request.original_content)} chars")
        
        # Create new document
        doc = Document()
        
        # Apply anonymization
        anonymized_content = DocumentProcessor.apply_anonymization(
            request.original_content, 
            request.entities
        )
        
        print(f"📝 Anonymized content length: {len(anonymized_content)} chars")
        print(f"📝 Anonymized preview: {anonymized_content[:200]}...")
        
        # Add content to document (split by lines)
        lines = anonymized_content.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip())
            else:
                doc.add_paragraph("")  # Empty line
        
        # Save to memory
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        print(f"✅ Document generated successfully, size: {len(doc_io.getvalue())} bytes")
        
        # Return as streaming response
        return StreamingResponse(
            io.BytesIO(doc_io.getvalue()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={request.filename}"}
        )
    
    except Exception as e:
        print(f"❌ Error generating document: {str(e)}")
        print(f"❌ Error type: {type(e).__name__}")
        import traceback
        print(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")

# Include router
app.include_router(api_router)

# CORRECTED: CORS middleware - Allow both ports
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",  # React dev server
        "http://127.0.0.1:3000",
        "http://localhost:3001",  # Backup port
        "http://127.0.0.1:3001",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting Anonymiseur Juridique RGPD v3.0 Backend")
    print("📍 Backend URL: http://localhost:8080")
    print(f"🧠 spaCy Available: {nlp is not None}")
    print(f"🔗 Ollama Support: {httpx is not None}")
    print("=" * 50)
    uvicorn.run(app, host="127.0.0.1", port=8080, reload=True)